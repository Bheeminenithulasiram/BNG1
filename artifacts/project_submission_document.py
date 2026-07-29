import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy/Dark Blue
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x40, 0x6E, 0x8D) # Secondary Accent
    return p

def add_body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33) # Charcoal
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    r2 = p.add_run(text)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    
    # We use a single-cell table for code background shading
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(5.7)
    set_cell_background(cell, "F2F2F2") # Light grey fill
    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
    
    # Remove borders
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="CCCCCC"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    cp = cell.paragraphs[0]
    cp.paragraph_format.space_after = Pt(0)
    cp.paragraph_format.line_spacing = 1.0
    run = cp.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

def main(output_path):
    doc = Document()
    
    # Set document margins (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # --- Title Page / Header ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    run_title = title_p.add_run("BrandGen — AI Brand Name Generator")
    run_title.font.name = 'Calibri Light'
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Navy
    
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    run_sub = sub_p.add_run("Project Showcasing & Submission Document")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(13)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12) # Spacing
    
    # --- 1. Project Name ---
    add_heading_1(doc, "1. Project Name")
    add_body(doc, "BrandGen — AI Brand Name Generator", bold=True)
    
    # --- 2. Problem Statement ---
    add_heading_1(doc, "2. Problem Statement")
    add_body(doc, "Finding a brand name for a new startup or project is a tedious process. Founders frequently face two major obstacles:")
    add_bullet(doc, "Most dictionary words or simple combinations are already registered on .com.", bold_prefix="Domain Congestion: ")
    add_bullet(doc, "Even if a domain is available, matching handles on major social platforms (Instagram, X/Twitter, LinkedIn) are often taken.", bold_prefix="Social Handle Fragmentation: ")
    add_body(doc, "Generic AI models suggest names but do not verify their domain status. Manually checking availability for 30+ candidates takes hours. BrandGen automates this by combining creative generative AI naming with instant multi-channel availability lookups.")
    
    # --- 3. Detailed Functionality Document ---
    add_heading_1(doc, "3. Detailed Functionality Document")
    add_body(doc, "BrandGen guides users through a structured name-discovery workflow:")
    add_bullet(doc, "The user describes their business (e.g., 'vintage film camera store'), selects an industry category, and adds style hints (e.g., 'minimalist').", bold_prefix="User Input: ")
    add_bullet(doc, "The Express backend queries the Groq API (LLaMA 3.3 70B model) to generate highly relevant, professional brand name candidates.", bold_prefix="AI Generation: ")
    add_bullet(doc, "The system evaluates .com availability and filters them into a balanced 80% available and 20% taken list of 18 candidates.", bold_prefix="Availability Harvesting: ")
    add_bullet(doc, "The user sees the 18 brand names displayed as interactive chips. Clicking a chip loads the Real-Time Verification Panel.", bold_prefix="Interactive Dashboard: ")
    add_bullet(doc, "The application checks and displays availability status across the .com registry (via RDAP/DNS) and handles on Instagram, X/Twitter, and LinkedIn.", bold_prefix="Real-Time Channel Verification: ")
    add_bullet(doc, "Users can test any custom name directly in the 'Check Custom Name' tab.", bold_prefix="Direct Search: ")
    
    # --- 4. Use Cases ---
    add_heading_1(doc, "4. Use Cases")
    add_bullet(doc, "Rapidly discover and lock down names with matched domains and social handles before launch.", bold_prefix="Startup Founders & Solopreneurs: ")
    add_bullet(doc, "Generate large lists of catchy, pre-validated compound brand names for client review.", bold_prefix="Domain Flippers & Brand Agencies: ")
    add_bullet(doc, "Brainstorm naming projects with clients in real-time, instantly weeding out taken names.", bold_prefix="Creative Agencies: ")
    
    # --- 5. Key Features ---
    add_heading_1(doc, "5. Key Features")
    add_bullet(doc, "Generates a realistic distribution of name candidates (14 likely available compound names and 4 common short taken names) to mirror real-world market dynamics.", bold_prefix="80/20 Domain Availability Mix: ")
    add_bullet(doc, "Combines official Verisign RDAP registry calls with local Node.js DNS lookups and Google DNS-over-HTTPS (DoH) fallbacks.", bold_prefix="Multi-Layered Domain Check: ")
    add_bullet(doc, "Utilizes public endpoints (Twitter OEmbed, Instagram HTML scrapers, LinkedIn profile crawlers) for real-time handle verification.", bold_prefix="Multi-Channel Social Check: ")
    add_bullet(doc, "Restricts traffic using express-rate-limit (20 req/min on AI generation), applies strict HTTP security headers via helmet, enforces request body limits, and configures CORS origin whitelist options.", bold_prefix="Production-Grade API Hardening: ")
    add_bullet(doc, "Features clean user input forms, scroll-to-section navigation, custom selectors, and a dark/light responsive design system.", bold_prefix="Responsive Single-Page Layout: ")
    
    # --- 6. Technology Stack Used ---
    add_heading_1(doc, "6. Technology Stack Used")
    add_heading_2(doc, "Frontend Stack")
    add_bullet(doc, "React 19 (via Vite 7)")
    add_bullet(doc, "Tailwind CSS v4 (Styling)")
    add_bullet(doc, "Radix UI & Lucide Icons (UI components & indicators)")
    add_bullet(doc, "TanStack React Query & React Hook Form (State management & validations)")
    add_bullet(doc, "Wouter (Routing)")
    
    add_heading_2(doc, "Backend Stack")
    add_bullet(doc, "Node.js 20+ (Express 5 web framework)")
    add_bullet(doc, "Pino (Structured Logger)")
    add_bullet(doc, "Helmet, express-rate-limit, CORS (Security & rate control)")
    
    add_heading_2(doc, "Infrastructure & Dev Tools")
    add_bullet(doc, "Groq API (LLaMA 3.3 70B model)")
    add_bullet(doc, "esbuild (ESM bundler)")
    add_bullet(doc, "Vercel (Hosting & Serverless Functions)")
    
    # --- 7. Architecture/Workflow ---
    add_heading_1(doc, "7. Architecture/Workflow")
    add_body(doc, "The application follows a client-server architecture. The user submits a form on the React frontend, which sends a POST request to the Express API backend hosted as a Serverless function. The backend orchestrates the call to the LLaMA 3.3 LLM via Groq, performs parallel domain check routines (RDAP/DNS/DoH), filters them to a matching 80/20 availability ratio, and returns the list. On select, the frontend queries social availability in parallel to display handle statuses.")
    
    # --- 8. API Documentation ---
    add_heading_1(doc, "8. API Documentation")
    
    add_heading_2(doc, "I. Generate Brand Names")
    add_body(doc, "Endpoint: POST /api/brands/generate", bold=True)
    add_body(doc, "Content-Type: application/json", italic=True)
    add_body(doc, "Request Body Format:")
    add_code_block(doc, '{\n  "description": "vintage film cameras and lenses shop",\n  "category": "ecommerce",\n  "keywords": "minimalist, retro"\n}')
    add_body(doc, "Response Body Format:")
    add_code_block(doc, '[\n  {\n    "name": "RetroHaven",\n    "tagline": "Make it instantly memorable",\n    "suggestedDomain": "retrohaven.com"\n  }\n]')
    
    add_heading_2(doc, "II. Check Handle & Domain Availability")
    add_body(doc, "Endpoint: POST /api/brands/availability", bold=True)
    add_body(doc, "Content-Type: application/json", italic=True)
    add_body(doc, "Request Body Format:")
    add_code_block(doc, '{\n  "name": "RetroHaven",\n  "domain": "retrohaven.com"\n}')
    add_body(doc, "Response Body Format:")
    add_code_block(doc, '{\n  "domain": { "name": "retrohaven.com", "status": "available" },\n  "social": {\n    "instagram": "available",\n    "twitter": "taken",\n    "linkedin": "available"\n  }\n}')

    # --- 9. GitHub/Source Code Repository Link ---
    add_heading_1(doc, "9. GitHub/Source Code Repository Link")
    add_body(doc, "https://github.com/Bheeminenithulasiram/BNG1")
    
    # --- 10. Live/Deployed Application Link ---
    add_heading_1(doc, "10. Live/Deployed Application Link")
    add_body(doc, "https://bng1.vercel.app")
    
    # --- 11. Screenshots or Demo Video ---
    add_heading_1(doc, "11. Screenshots or Demo Video")
    add_body(doc, "The following screenshots showcase the interface and validation functionality of the BrandGen web application:")
    
    import glob
    import os
    
    artifact_dir = r"C:\Users\SARANYA ATREYAPURAPU\.gemini\antigravity-ide\brain\b1175263-549a-40b3-afbb-df0922cb5fc3"
    
    landing_images = glob.glob(os.path.join(artifact_dir, "landing_page_*.png"))
    generated_images = glob.glob(os.path.join(artifact_dir, "generated_brands_*.png"))
    details_images = glob.glob(os.path.join(artifact_dir, "brand_details_*.png"))
    
    if landing_images:
        add_body(doc, "Figure 1: BrandGen Generator Form (Landing Page)", bold=True, space_after=2)
        doc.add_picture(landing_images[0], width=Inches(5.5))
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
    if generated_images:
        add_body(doc, "Figure 2: Generated Brand Name Candidates Grid (18 Suggestions in 80/20 Mix)", bold=True, space_after=2)
        doc.add_picture(generated_images[0], width=Inches(5.5))
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
    if details_images:
        add_body(doc, "Figure 3: Multi-Channel Availability Validation Panel for Chosen Brand", bold=True, space_after=2)
        doc.add_picture(details_images[0], width=Inches(5.5))
        doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
    add_body(doc, "Walkthrough Video: The complete user flow is recorded in the 'brandgen_demo_flow_*.webp' video file included in the project directory.")
    
    # --- 12. Your Contribution to the Project ---
    add_heading_1(doc, "12. Your Contribution to the Project")
    add_bullet(doc, "Replaced the baseline 'all available' filter constraint with a balanced 80% available / 20% taken domain availability ratio model, matching real-world expectations.", bold_prefix="Feature Design: ")
    add_bullet(doc, "Structured prompt guidelines for LLaMA 3.3 to yield a diverse blend of creative compound phrases (high-availability) and short root terms (taken).", bold_prefix="AI Prompt Engineering: ")
    add_bullet(doc, "Developed a fast, two-round async candidate checking loop optimized to complete checks within Serverless latency thresholds.", bold_prefix="Multi-Round Harvester Loop: ")
    add_bullet(doc, "Implemented Helmet middleware headers, Express IP rate limiters, body-size restrictions, and customizable CORS options for secure server deployments.", bold_prefix="API Security Hardening: ")
    add_bullet(doc, "Resolved critical Vercel compilation bugs by refactoring ES module imports, cleaning root dependencies, and defining native build commands.", bold_prefix="Vercel Integration Fixes: ")

    doc.save(output_path)
    print(f"DOCX created successfully at: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_docx.py <output_path>")
        sys.exit(1)
    main(sys.argv[1])
